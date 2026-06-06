"""
tools/export_tools.py
──────────────────────
Converts a Markdown proposal string into professional Word (.docx)
and PDF documents for download.

Libraries used (both open-source, no external tools needed)
────────────────────────────────────────────────────────────
  • python-docx  — creates .docx with proper heading hierarchy,
                   paragraph styles, and a reference list
  • reportlab    — creates .pdf with a clean single-column academic layout

TUTORIAL NOTE
─────────────
The export pipeline is:

  proposal_markdown (str)
          │
          ├──► parse_proposal_sections()   split on ## headings
          │
          ├──► build_docx()                → bytes (Word document)
          │
          └──► build_pdf()                 → bytes (PDF document)

Both functions return raw bytes so Streamlit's st.download_button()
can serve them directly without writing to disk.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple

# ── Section parser ────────────────────────────────────────────────────────────

def parse_proposal_sections(markdown: str) -> List[Tuple[int, str, str]]:
    """
    Split Markdown into (level, heading, body) tuples.

    Returns a list like:
      [(1, "Research Proposal", ""),
       (2, "Abstract", "Lorem ipsum..."),
       (2, "Introduction", "..."),
       ...]

    Handles up to heading level 3.
    """
    sections: List[Tuple[int, str, str]] = []
    current_level = 0
    current_heading = ""
    current_body_lines: List[str] = []

    for line in markdown.splitlines():
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if m:
            if current_heading or current_body_lines:
                sections.append((
                    current_level,
                    current_heading,
                    "\n".join(current_body_lines).strip(),
                ))
            current_level = len(m.group(1))
            current_heading = m.group(2).strip()
            current_body_lines = []
        else:
            current_body_lines.append(line)

    if current_heading or current_body_lines:
        sections.append((
            current_level,
            current_heading,
            "\n".join(current_body_lines).strip(),
        ))

    return sections


def _strip_inline_md(text: str) -> str:
    """Remove bold/italic markers but preserve citation brackets."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    return text


# ── DOCX Export ───────────────────────────────────────────────────────────────

def build_docx(
    proposal_markdown: str,
    references: List[Dict],
    metadata: Optional[Dict] = None,
) -> bytes:
    """
    Build a Word document from proposal Markdown.

    Styling choices (following standard academic proposal conventions):
    - Calibri 12pt body text (universally readable)
    - 1.5 line spacing
    - Headings use built-in Word styles (Heading 1/2/3) for navigation
    - References section uses 10pt Calibri with hanging indent
    - Margins: 2.5 cm all sides

    Parameters
    ----------
    proposal_markdown : full proposal as Markdown string
    references        : list of reference dicts (must have 'apa' and 'ref_num')
    metadata          : optional dict with 'author', 'institution', 'date'

    Returns
    -------
    Raw bytes of the .docx file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("pip install python-docx")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── Default paragraph style ───────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    para_fmt = style.paragraph_format
    para_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para_fmt.space_after = Pt(6)

    # ── Header meta ───────────────────────────────────────────
    if metadata:
        hdr = doc.add_paragraph()
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr.paragraph_format.space_after = Pt(2)
        run = hdr.add_run(metadata.get("institution", ""))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        date_str = metadata.get("date", datetime.today().strftime("%B %Y"))
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.add_run(date_str).font.size = Pt(10)
        doc.add_paragraph()  # spacer

    # ── Parse and render sections ─────────────────────────────
    sections = parse_proposal_sections(proposal_markdown)

    for level, heading, body in sections:
        # Headings
        if level == 1:
            h = doc.add_heading(heading, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.runs[0].font.size = Pt(16)
        elif level == 2:
            h = doc.add_heading(heading, level=2)
            h.runs[0].font.size = Pt(13)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        elif level == 3:
            h = doc.add_heading(heading, level=3)
            h.runs[0].font.size = Pt(12)
        else:
            # Treat as bold paragraph if level 0 (preamble text)
            if heading:
                p = doc.add_paragraph()
                p.add_run(heading).bold = True

        # Body text — split on blank lines → paragraphs
        if body:
            paragraphs = re.split(r"\n{2,}", body)
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # Ordered / unordered list items
                if re.match(r"^\d+\.\s", para_text) or para_text.startswith("- "):
                    for item in para_text.splitlines():
                        item = item.strip().lstrip("0123456789.-").strip()
                        if item:
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(_strip_inline_md(item)).font.size = Pt(12)
                    continue

                # Markdown table → plain text (simplified)
                if para_text.startswith("|"):
                    rows = [
                        [cell.strip() for cell in row.strip("|").split("|")]
                        for row in para_text.splitlines()
                        if not re.match(r"^\|[-| :]+\|$", row.strip())
                    ]
                    if rows:
                        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        tbl.style = "Table Grid"
                        for r_i, row in enumerate(rows):
                            for c_i, cell in enumerate(row):
                                tbl.cell(r_i, c_i).text = cell
                    continue

                # Regular paragraph
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(6)

                # Render bold **text** inline
                parts = re.split(r"(\*\*.+?\*\*)", para_text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(12)
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(12)

    # ── References section (if not already in markdown) ───────
    already_has_refs = any(
        "reference" in h.lower() for _, h, _ in sections
    )
    if references and not already_has_refs:
        doc.add_heading("References", level=2)
        for ref in references:
            apa = ref.get("apa", ref.get("title", "Unknown"))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)  # hanging indent
            p.paragraph_format.space_after = Pt(4)
            p.add_run(f"[{ref.get('ref_num', '?')}] ").bold = True
            p.add_run(apa).font.size = Pt(10)

    # ── Serialise to bytes ────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF Export ────────────────────────────────────────────────────────────────

def build_pdf(
    proposal_markdown: str,
    references: List[Dict],
    metadata: Optional[Dict] = None,
) -> bytes:
    """
    Build a PDF from proposal Markdown using ReportLab.

    Layout:
    - A4 page, 2.5 cm margins
    - Title: 18pt bold centred
    - Section headings: 13pt bold, dark blue
    - Body text: 11pt, justified, 1.3× line spacing
    - References: 9pt, hanging indent

    Parameters / returns same as build_docx().
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, ListFlowable, ListItem,
        )
        from reportlab.platypus import HRFlowable
    except ImportError:
        raise ImportError("pip install reportlab")

    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=3.0 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=metadata.get("title", "Research Proposal") if metadata else "Research Proposal",
        author=metadata.get("author", "") if metadata else "",
    )

    # ── Styles ────────────────────────────────────────────────
    styles = getSampleStyleSheet()

    DARK_BLUE = colors.HexColor("#1F3964")
    MID_GREY  = colors.HexColor("#444444")

    title_style = ParagraphStyle(
        "ProposalTitle",
        parent=styles["Title"],
        fontSize=18,
        leading=24,
        alignment=TA_CENTER,
        textColor=DARK_BLUE,
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=10,
        alignment=TA_CENTER,
        textColor=MID_GREY,
        spaceAfter=4,
    )
    h2_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontSize=13,
        leading=17,
        textColor=DARK_BLUE,
        spaceBefore=14,
        spaceAfter=4,
        fontName="Helvetica-Bold",
    )
    h3_style = ParagraphStyle(
        "SubHeading",
        parent=styles["Heading3"],
        fontSize=11,
        leading=14,
        textColor=MID_GREY,
        spaceBefore=8,
        spaceAfter=3,
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "ProposalBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=styles["Normal"],
        fontSize=11,
        leading=15,
        leftIndent=18,
        bulletIndent=4,
        spaceAfter=3,
    )
    ref_style = ParagraphStyle(
        "Reference",
        parent=styles["Normal"],
        fontSize=9,
        leading=13,
        leftIndent=24,
        firstLineIndent=-24,
        spaceAfter=4,
        textColor=MID_GREY,
    )

    story = []

    # ── Header metadata ───────────────────────────────────────
    if metadata:
        if metadata.get("institution"):
            story.append(Paragraph(metadata["institution"], meta_style))
        date_str = metadata.get("date", datetime.today().strftime("%B %Y"))
        story.append(Paragraph(date_str, meta_style))
        story.append(Spacer(1, 0.3 * cm))

    # ── Parse and render sections ─────────────────────────────
    sections = parse_proposal_sections(proposal_markdown)
    already_has_refs = any("reference" in h.lower() for _, h, _ in sections)

    for level, heading, body in sections:
        if level == 1:
            story.append(Paragraph(_escape_rl(heading), title_style))
            story.append(HRFlowable(width="100%", thickness=1, color=DARK_BLUE, spaceAfter=8))
        elif level == 2:
            story.append(Paragraph(_escape_rl(heading), h2_style))
        elif level == 3:
            story.append(Paragraph(_escape_rl(heading), h3_style))
        elif heading:
            story.append(Paragraph(f"<b>{_escape_rl(heading)}</b>", body_style))

        if body:
            paragraphs = re.split(r"\n{2,}", body)
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # List items
                if re.match(r"^[-*]\s", para_text) or re.match(r"^\d+\.\s", para_text):
                    for item in para_text.splitlines():
                        item = item.strip().lstrip("0123456789.-*").strip()
                        if item:
                            story.append(
                                Paragraph(f"• {_md_to_rl(item)}", bullet_style)
                            )
                    continue

                # Markdown table → ReportLab Table
                if para_text.startswith("|"):
                    rows = [
                        [cell.strip() for cell in row.strip("|").split("|")]
                        for row in para_text.splitlines()
                        if not re.match(r"^\|[-| :]+\|$", row.strip())
                    ]
                    if rows:
                        col_count = max(len(r) for r in rows)
                        # Pad short rows
                        rows = [r + [""] * (col_count - len(r)) for r in rows]
                        tbl = Table(rows, hAlign="LEFT")
                        tbl.setStyle(TableStyle([
                            ("BACKGROUND", (0, 0), (-1, 0), DARK_BLUE),
                            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE",   (0, 0), (-1, -1), 9),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4F8")]),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
                            ("TOPPADDING", (0, 0), (-1, -1), 4),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ]))
                        story.append(tbl)
                        story.append(Spacer(1, 0.2 * cm))
                    continue

                story.append(Paragraph(_md_to_rl(para_text), body_style))

    # ── References ────────────────────────────────────────────
    if references and not already_has_refs:
        story.append(Paragraph("References", h2_style))
        story.append(HRFlowable(width="60%", thickness=0.5, color=colors.lightgrey, spaceAfter=4))
        for ref in references:
            apa = ref.get("apa", ref.get("title", "Unknown"))
            num = ref.get("ref_num", "?")
            story.append(
                Paragraph(f"<b>[{num}]</b> {_escape_rl(apa)}", ref_style)
            )

    # ── Page footer ───────────────────────────────────────────
    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(MID_GREY)
        canvas.drawCentredString(
            A4[0] / 2, 1.5 * cm,
            f"Page {doc.page} — Generated by Agentic Research Assistant"
        )
        canvas.restoreState()

    doc_rl.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


# ── ReportLab text helpers ─────────────────────────────────────────────────────

def _escape_rl(text: str) -> str:
    """Escape characters that ReportLab treats as XML/HTML."""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


def _md_to_rl(text: str) -> str:
    """Convert inline Markdown (bold, italic, citations) to ReportLab XML."""
    # Bold **text**
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Italic *text*
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    # Citation brackets [Author, Year] — keep as-is, just escape surrounding text
    # (already handled by bold/italic patterns above)
    return _escape_rl_partial(text)


def _escape_rl_partial(text: str) -> str:
    """Escape & and bare < / > that aren't already inside XML tags."""
    # Only escape & not already escaped
    text = re.sub(r"&(?!amp;|lt;|gt;|#)", "&amp;", text)
    # Escape < and > that are not part of RL tags
    text = re.sub(r"<(?!/?(?:b|i|u|br|super|sub|font|para)\b)", "&lt;", text)
    return text
