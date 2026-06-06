"""
tools/prisma_report.py
──────────────────────
Generate a PRISMA 2020-compliant systematic review manuscript from a
completed SystematicReviewState dict.

Two output formats:
  DOCX  — via python-docx (already in requirements); Word/Google Docs compatible
  PDF   — via reportlab (already in requirements); pure-Python, no LibreOffice

PRISMA 2020 section order:
  Title · Abstract · Introduction · Methods (Search Strategy, Eligibility Criteria,
  Data Extraction, Quality Assessment) · Results (Study Selection, Study
  Characteristics, Evidence Table, Narrative Synthesis) · Discussion (Summary,
  Research Gaps, Limitations, Conclusions) · References
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any, Dict, List

logger = logging.getLogger(__name__)

_TODAY = datetime.today().strftime("%B %d, %Y")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def generate_prisma_docx(
    state: Dict[str, Any],
    author: str = "",
    institution: str = "",
) -> bytes:
    """Return .docx bytes of a PRISMA 2020-compliant manuscript."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("pip install python-docx")

    rq = state.get("research_question", "")
    flow = state.get("prisma_flow", {})
    evidence_table = state.get("evidence_table", [])
    included = state.get("included_papers", [])
    excluded = state.get("excluded_papers", [])
    narrative = state.get("narrative_synthesis", "")
    themes = state.get("key_themes", [])
    gaps = state.get("research_gaps", [])
    conclusion = state.get("conclusion", "")
    limitations = state.get("limitations", "")
    queries = state.get("search_queries", [])
    inc_criteria = state.get("inclusion_criteria", [])
    exc_criteria = state.get("exclusion_criteria", [])

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Title page ────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Systematic Review Report")
    run.bold = True
    run.font.size = Pt(18)

    if rq:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(rq)
        sub_run.italic = True
        sub_run.font.size = Pt(12)

    meta_lines = [s for s in [author, institution, _TODAY] if s]
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.add_run("\n".join(meta_lines)).font.size = Pt(11)

    doc.add_page_break()

    # ── Abstract ─────────────────────────────────────────────────────────────
    doc.add_heading("Abstract", level=1)
    abstract_text = (
        f"Background: This systematic review addresses the following research question: {rq}. "
        f"Methods: A systematic search was conducted across academic databases using {len(queries)} queries. "
        f"Results: {flow.get('identified', 0)} records were identified; "
        f"{flow.get('screened', 0)} were screened; "
        f"{flow.get('included', 0)} studies met all inclusion criteria. "
        f"Key themes: {'; '.join(themes[:3]) if themes else 'see main text'}. "
        f"Conclusions: {(conclusion or 'See main text.')[:250]}"
    )
    doc.add_paragraph(abstract_text)
    doc.add_heading("Keywords", level=3)
    kw = [w for w in rq.split() if len(w) > 4][:6]
    doc.add_paragraph(", ".join(kw))

    # ── Introduction ─────────────────────────────────────────────────────────
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This systematic review was conducted in accordance with the Preferred Reporting Items for "
        "Systematic Reviews and Meta-Analyses (PRISMA) 2020 guidelines. "
        f"The primary research question investigated was: {rq}"
    )
    if gaps:
        doc.add_paragraph("This review was motivated by the following gaps in existing literature: " +
                          "; ".join(gaps[:3]) + ".")

    # ── Methods ──────────────────────────────────────────────────────────────
    doc.add_heading("Methods", level=1)

    doc.add_heading("Search Strategy", level=2)
    doc.add_paragraph(
        f"A systematic search was conducted across arXiv, Semantic Scholar, CrossRef, and Google Scholar "
        f"using {len(queries)} distinct queries:"
    )
    for i, q in enumerate(queries, 1):
        doc.add_paragraph(f"{i}. {q}", style="List Number")

    doc.add_heading("Eligibility Criteria", level=2)
    doc.add_heading("Inclusion Criteria", level=3)
    for c in (inc_criteria or ["Papers relevant to the research question"]):
        doc.add_paragraph(c, style="List Bullet")
    doc.add_heading("Exclusion Criteria", level=3)
    for c in (exc_criteria or ["Papers clearly off-topic"]):
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("Data Extraction", level=2)
    doc.add_paragraph(
        "Data were extracted using a standardised form covering: study design, sample size, "
        "population, key findings, and quality rating. All extraction was performed by the automated agent."
    )

    doc.add_heading("Quality Assessment", level=2)
    doc.add_paragraph(
        "Study quality was rated High, Medium, or Low based on study design hierarchy "
        "(RCTs > cohort > cross-sectional > narrative review) and methodological transparency."
    )

    # ── Results ──────────────────────────────────────────────────────────────
    doc.add_heading("Results", level=1)

    doc.add_heading("Study Selection", level=2)
    doc.add_paragraph(
        f"Database searches identified {flow.get('identified', 0)} records. "
        f"After deduplication and title/abstract screening, {flow.get('screened', 0)} records were assessed. "
        f"Of these, {flow.get('included', 0)} studies met all inclusion criteria and were included. "
        f"{flow.get('excluded', 0)} records were excluded."
    )

    # PRISMA flow table
    doc.add_heading("PRISMA Flow", level=3)
    prisma_tbl = doc.add_table(rows=6, cols=2)
    prisma_tbl.style = "Table Grid"
    rows_data = [
        ("Stage", "Count"),
        ("Identified", str(flow.get("identified", 0))),
        ("Screened", str(flow.get("screened", 0))),
        ("Eligibility assessed", str(flow.get("eligibility", 0))),
        ("Included", str(flow.get("included", 0))),
        ("Excluded", str(flow.get("excluded", 0))),
    ]
    for i, (stage, count) in enumerate(rows_data):
        row = prisma_tbl.rows[i]
        row.cells[0].text = stage
        row.cells[1].text = count
        if i == 0:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()

    doc.add_heading("Study Characteristics", level=2)
    design_counts: Dict[str, int] = {}
    quality_counts: Dict[str, int] = {}
    for row in evidence_table:
        d = row.get("study_design", "Unknown")
        q = row.get("quality", "Medium")
        design_counts[d] = design_counts.get(d, 0) + 1
        quality_counts[q] = quality_counts.get(q, 0) + 1
    if design_counts:
        doc.add_paragraph("Study designs: " +
                          ", ".join(f"{d} (n={n})" for d, n in sorted(design_counts.items(), key=lambda x: -x[1])) + ".")
    if quality_counts:
        doc.add_paragraph("Quality distribution: " +
                          ", ".join(f"{q}: {n}" for q, n in sorted(quality_counts.items())) + ".")

    # Evidence table
    if evidence_table:
        doc.add_heading("Evidence Table", level=3)
        ev_tbl = doc.add_table(rows=1 + min(len(evidence_table), 20), cols=5)
        ev_tbl.style = "Table Grid"
        for cell, hdr in zip(ev_tbl.rows[0].cells, ["Citation", "Year", "Design", "Quality", "Key Finding"]):
            cell.text = hdr
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for i, ev in enumerate(evidence_table[:20], 1):
            row = ev_tbl.rows[i]
            row.cells[0].text = ev.get("citation_key", "")
            row.cells[1].text = str(ev.get("year", "n.d."))
            row.cells[2].text = ev.get("study_design", "")
            row.cells[3].text = ev.get("quality", "")
            row.cells[4].text = ev.get("key_finding", "")[:100]
        doc.add_paragraph()

    if themes:
        doc.add_heading("Key Themes", level=2)
        for theme in themes:
            doc.add_paragraph(theme, style="List Bullet")

    doc.add_heading("Narrative Synthesis", level=2)
    for para_text in narrative.split("\n\n"):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

    # ── Discussion ────────────────────────────────────────────────────────────
    doc.add_heading("Discussion", level=1)

    doc.add_heading("Summary of Evidence", level=2)
    doc.add_paragraph(conclusion or "See narrative synthesis above.")

    if gaps:
        doc.add_heading("Research Gaps", level=2)
        for g in gaps:
            doc.add_paragraph(g, style="List Bullet")

    doc.add_heading("Limitations", level=2)
    doc.add_paragraph(limitations or (
        "This review is subject to limitations inherent in automated systematic reviews, "
        "including potential incomplete database coverage and reliance on title/abstract screening."
    ))

    doc.add_heading("Conclusions", level=2)
    doc.add_paragraph(conclusion or "Conclusions are presented in the narrative synthesis section.")

    # ── References ────────────────────────────────────────────────────────────
    doc.add_heading("References", level=1)
    seen: set = set()
    ref_n = 1
    for paper in included:
        ck = paper.get("citation_key", "")
        if ck in seen:
            continue
        seen.add(ck)
        authors = "; ".join(paper.get("authors", ["Unknown"])[:3])
        yr = paper.get("year", "n.d.")
        title_str = paper.get("title", "")
        journal = paper.get("journal", "Preprint")
        doi = paper.get("doi", "")
        url = paper.get("url", "")
        link = f"https://doi.org/{doi}" if doi else url
        doc.add_paragraph(f"{ref_n}. {authors} ({yr}). {title_str}. {journal}. {link}",
                          style="List Number")
        ref_n += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def generate_prisma_pdf(
    state: Dict[str, Any],
    author: str = "",
    institution: str = "",
) -> bytes:
    """Return .pdf bytes of a PRISMA 2020-compliant manuscript (reportlab)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, PageBreak,
        )
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    except ImportError:
        raise ImportError("pip install reportlab")

    rq = state.get("research_question", "")
    flow = state.get("prisma_flow", {})
    evidence_table = state.get("evidence_table", [])
    included = state.get("included_papers", [])
    narrative = state.get("narrative_synthesis", "")
    themes = state.get("key_themes", [])
    gaps = state.get("research_gaps", [])
    conclusion = state.get("conclusion", "")
    limitations = state.get("limitations", "")
    queries = state.get("search_queries", [])
    inc_criteria = state.get("inclusion_criteria", [])
    exc_criteria = state.get("exclusion_criteria", [])

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2.5*cm, rightMargin=2.5*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    accent = HexColor("#1a6496")

    H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=15, textColor=accent,
                        spaceBefore=14, spaceAfter=6)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, textColor=accent,
                        spaceBefore=10, spaceAfter=4)
    H3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=10, textColor=black,
                        spaceBefore=8, spaceAfter=3)
    BODY = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14,
                          alignment=TA_JUSTIFY, spaceAfter=5)
    BULLET = ParagraphStyle("Bullet", parent=BODY, leftIndent=18, bulletIndent=8,
                            spaceBefore=2, spaceAfter=2)
    TITLE_S = ParagraphStyle("TitleS", parent=styles["Title"], fontSize=18, leading=22,
                              alignment=TA_CENTER, textColor=accent)
    CENTER = ParagraphStyle("Center", parent=BODY, alignment=TA_CENTER)

    story = []

    def h1(t): story.append(Paragraph(t, H1))
    def h2(t): story.append(Paragraph(t, H2))
    def h3(t): story.append(Paragraph(t, H3))
    def para(t): story.append(Paragraph(t or " ", BODY))
    def bullet(t): story.append(Paragraph(f"• {t}", BULLET))
    def space(n=6): story.append(Spacer(1, n))

    # Title page
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph("Systematic Review Report", TITLE_S))
    if rq:
        story.append(Paragraph(f"<i>{rq}</i>",
                               ParagraphStyle("Sub", parent=BODY, alignment=TA_CENTER, fontSize=11)))
    space(8)
    meta = [s for s in [author, institution, _TODAY] if s]
    story.append(Paragraph("<br/>".join(meta), CENTER))
    story.append(PageBreak())

    # Abstract
    h1("Abstract")
    para(
        f"<b>Background:</b> {rq}. "
        f"<b>Methods:</b> Systematic search using {len(queries)} queries across multiple databases. "
        f"<b>Results:</b> {flow.get('identified', 0)} records identified; "
        f"{flow.get('included', 0)} studies included. "
        f"Key themes: {'; '.join(themes[:3]) if themes else 'see main text'}. "
        f"<b>Conclusions:</b> {(conclusion or 'See main text.')[:250]}"
    )
    kw = [w for w in rq.split() if len(w) > 4][:6]
    if kw:
        h3("Keywords")
        para(", ".join(kw))
    space(10)

    # Introduction
    h1("Introduction")
    para(
        "This systematic review was conducted in accordance with PRISMA 2020 guidelines. "
        f"The primary research question was: <i>{rq}</i>"
    )
    if gaps:
        para("Motivated by gaps: " + "; ".join(gaps[:3]) + ".")
    space()

    # Methods
    h1("Methods")
    h2("Search Strategy")
    para(f"Searches across arXiv, Semantic Scholar, CrossRef, and Google Scholar using {len(queries)} queries:")
    for i, q in enumerate(queries, 1):
        bullet(f"{i}. {q}")
    h2("Eligibility Criteria")
    h3("Inclusion Criteria")
    for c in (inc_criteria or ["Relevant to the research question"]):
        bullet(c)
    h3("Exclusion Criteria")
    for c in (exc_criteria or ["Clearly off-topic"]):
        bullet(c)
    h2("Data Extraction")
    para("Standardised extraction: study design, sample size, key finding, quality rating.")
    h2("Quality Assessment")
    para("Quality rated High/Medium/Low per study design and methodological transparency.")
    space()

    # Results
    h1("Results")
    h2("Study Selection")
    para(
        f"{flow.get('identified', 0)} records identified. "
        f"{flow.get('screened', 0)} screened. "
        f"{flow.get('included', 0)} studies included; {flow.get('excluded', 0)} excluded."
    )

    # PRISMA table
    h3("PRISMA Flow")
    flow_rows = [
        ["Stage", "Count"],
        ["Identified", str(flow.get("identified", 0))],
        ["Screened", str(flow.get("screened", 0))],
        ["Eligibility assessed", str(flow.get("eligibility", 0))],
        ["Included", str(flow.get("included", 0))],
        ["Excluded", str(flow.get("excluded", 0))],
    ]
    ft = Table(flow_rows, colWidths=[11*cm, 4*cm])
    ft.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), accent),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, black),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f0f4f8"), white]),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(ft)
    space(10)

    # Study characteristics
    h2("Study Characteristics")
    design_counts: Dict[str, int] = {}
    quality_counts: Dict[str, int] = {}
    for row in evidence_table:
        d = row.get("study_design", "Unknown")
        q = row.get("quality", "Medium")
        design_counts[d] = design_counts.get(d, 0) + 1
        quality_counts[q] = quality_counts.get(q, 0) + 1
    if design_counts:
        para("Study designs: " +
             ", ".join(f"{d} (n={n})" for d, n in sorted(design_counts.items(), key=lambda x: -x[1])) + ".")
    if quality_counts:
        para("Quality: " + ", ".join(f"{q}: {n}" for q, n in sorted(quality_counts.items())) + ".")

    # Evidence table
    if evidence_table:
        h3("Evidence Table (top 15 studies)")
        ev_rows = [["Citation", "Year", "Design", "Quality", "Key Finding"]]
        for ev in evidence_table[:15]:
            ev_rows.append([
                ev.get("citation_key", "")[:18],
                str(ev.get("year", "n.d.")),
                ev.get("study_design", "")[:14],
                ev.get("quality", ""),
                ev.get("key_finding", "")[:55],
            ])
        et = Table(ev_rows, colWidths=[3.2*cm, 1.4*cm, 3*cm, 2*cm, 7.4*cm])
        et.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), accent),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, black),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f0f4f8"), white]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(et)
        space(10)

    if themes:
        h2("Key Themes")
        for t in themes:
            bullet(t)

    h2("Narrative Synthesis")
    for chunk in narrative.split("\n\n"):
        if chunk.strip():
            para(chunk.strip())

    # Discussion
    h1("Discussion")
    h2("Summary of Evidence")
    para(conclusion or "See narrative synthesis above.")
    if gaps:
        h2("Research Gaps")
        for g in gaps:
            bullet(g)
    h2("Limitations")
    para(limitations or (
        "This review is subject to limitations inherent in automated systematic reviews, "
        "including potential incomplete database coverage and reliance on title/abstract screening."
    ))
    h2("Conclusions")
    para(conclusion or "Conclusions are presented in the narrative synthesis section.")

    # References
    h1("References")
    seen: set = set()
    ref_n = 1
    for paper in included:
        ck = paper.get("citation_key", "")
        if ck in seen:
            continue
        seen.add(ck)
        authors = "; ".join(paper.get("authors", ["Unknown"])[:3])
        yr = paper.get("year", "n.d.")
        title_str = paper.get("title", "")
        journal = paper.get("journal", "Preprint")
        doi = paper.get("doi", "")
        url = paper.get("url", "")
        link = f"https://doi.org/{doi}" if doi else url
        para(f"{ref_n}. {authors} ({yr}). {title_str}. <i>{journal}</i>. {link}")
        ref_n += 1

    doc.build(story)
    return buf.getvalue()
