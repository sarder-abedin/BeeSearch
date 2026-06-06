"""
tools/document_tools.py
────────────────────────
Handles ingestion and pre-processing of user-uploaded documents.

Supported formats (standard parser)
─────────────────────────────────────
  • PDF  — pdfplumber (tables + text, layout-aware)
  • DOCX — python-docx
  • TXT / MD — plain text

Use get_processor(use_docling=True) to get a DoclingProcessor that
additionally handles PPTX, XLSX, HTML, images, and provides OCR.

TUTORIAL NOTE — Vectorless RAG
───────────────────────────────
In vectorless RAG mode documents are injected directly into the LLM's context
window — no embedding or vector store is needed.  Because we only ever inject
up to `max_raw_chars` characters per document, we can stop reading pages early
once we have enough text, saving significant memory for large PDFs.

Setting max_raw_chars to 0 disables the cap (reads the entire file).
"""

from __future__ import annotations

import gc
import hashlib
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import IO, List, Optional, Union

logger = logging.getLogger(__name__)


# ── Data Classes ──────────────────────────────────────────────────────────────

@dataclass
class DocumentChunk:
    """A single piece of a document, ready for embedding."""
    chunk_id: str          # Stable hash of content
    doc_id: str            # Parent document identifier
    doc_name: str          # Human-readable filename
    page_num: int          # Source page (0-based; -1 if unavailable)
    chunk_index: int       # Position within the document
    text: str              # Raw extracted text
    metadata: dict = field(default_factory=dict)


@dataclass
class ProcessedDocument:
    """All chunks extracted from one uploaded file."""
    doc_id: str
    filename: str
    file_type: str
    total_pages: int
    total_chunks: int
    chunks: List[DocumentChunk]
    raw_text: str          # Full concatenated text (for summary prompts)
    content_md5: str = ""  # MD5 of first 50 000 chars of raw_text for cache invalidation


# ── Utilities ─────────────────────────────────────────────────────────────────

def _stable_id(text: str) -> str:
    """MD5-based deterministic ID; short enough for display, unique enough for practice."""
    return hashlib.md5(text.encode()).hexdigest()[:12]


def _clean_text(text: str) -> str:
    """Remove control characters, normalise whitespace."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def _chunk_text(
    text: str,
    chunk_size: int = 800,
    overlap: int = 150,
) -> List[str]:
    """
    Sliding-window character chunker with sentence-boundary awareness.

    We try to break at sentence endings ('. ', '! ', '? ') rather than
    mid-sentence so each chunk is semantically complete.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks: List[str] = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        # Try to extend to the next sentence boundary within a 200-char window
        if end < len(text):
            search_end = min(end + 200, len(text))
            boundary = -1
            for sep in (". ", "! ", "? ", "\n\n"):
                idx = text.rfind(sep, start, search_end)
                if idx > boundary:
                    boundary = idx
            if boundary > start:
                end = boundary + 2  # include the separator

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap  # overlap keeps context across boundaries

    return chunks


# ── Main Processor ─────────────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Stateless helper that converts uploaded files into DocumentChunk lists.

    Usage
    -----
    processor = DocumentProcessor(chunk_size=800, overlap=150)
    doc = processor.process_file("/tmp/paper.pdf")
    print(f"Extracted {doc.total_chunks} chunks from {doc.filename}")
    """

    def __init__(
        self,
        chunk_size: int = 800,
        overlap: int = 150,
        max_raw_chars: int = 0,
        max_pages: int = 300,
    ):
        self.chunk_size = chunk_size
        self.overlap = overlap
        # Stop reading pages once accumulated raw text reaches this many chars.
        # 0 = read the entire file. Set to the vectorless-RAG context budget to
        # avoid holding a full 300-page PDF in memory when only ~20 000 chars
        # will actually be injected into the LLM.
        self.max_raw_chars = max_raw_chars
        # Hard page cap — prevents OOM on very large PDFs. 0 = no limit.
        self.max_pages = max_pages

    # ── Public API ────────────────────────────────────────────

    def process_file(
        self,
        file_path: Union[str, Path],
        file_obj: Optional[IO[bytes]] = None,
    ) -> ProcessedDocument:
        """
        Process a file on disk (or an in-memory bytes stream from Streamlit).

        Parameters
        ----------
        file_path : path to the file (used for extension detection + doc_id)
        file_obj  : optional IO object when the file comes from a web upload
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        doc_id = _stable_id(path.name)

        logger.info("Processing %s (type=%s)", path.name, ext)

        if ext == ".pdf":
            pages_text, total_pages = self._extract_pdf(path, file_obj)
        elif ext in (".docx", ".doc"):
            pages_text, total_pages = self._extract_docx(path, file_obj)
        elif ext in (".txt", ".md", ".rst"):
            pages_text, total_pages = self._extract_text(path, file_obj)
        else:
            raise ValueError(
                f"Unsupported file type '{ext}'. Accepted: PDF, DOCX, TXT, MD."
            )

        raw_text = "\n\n".join(pages_text)
        content_md5 = hashlib.md5(raw_text[:50000].encode()).hexdigest()
        chunks = self._build_chunks(pages_text, doc_id, path.name)

        return ProcessedDocument(
            doc_id=doc_id,
            filename=path.name,
            file_type=ext.lstrip(".").upper(),
            total_pages=total_pages,
            total_chunks=len(chunks),
            chunks=chunks,
            raw_text=raw_text,
            content_md5=content_md5,
        )

    def process_raw_text(self, text: str, name: str = "pasted_text") -> ProcessedDocument:
        """Convenience method for plain text provided directly in the UI."""
        doc_id = _stable_id(text[:200])
        content_md5 = hashlib.md5(text[:50000].encode()).hexdigest()
        pages_text = [text]
        chunks = self._build_chunks(pages_text, doc_id, name)
        return ProcessedDocument(
            doc_id=doc_id,
            filename=name,
            file_type="TXT",
            total_pages=1,
            total_chunks=len(chunks),
            chunks=chunks,
            raw_text=text,
            content_md5=content_md5,
        )

    # ── Extraction Backends ───────────────────────────────────

    def _extract_pdf(
        self, path: Path, file_obj: Optional[IO[bytes]]
    ) -> tuple[List[str], int]:
        """Use pdfplumber for robust PDF extraction (handles tables too)."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("Install pdfplumber: pip install pdfplumber")

        source = file_obj if file_obj is not None else str(path)
        pages_text: List[str] = []

        total_chars = 0
        try:
            with pdfplumber.open(source) as pdf:
                total_pdf_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages):
                    # Hard page cap — prevents OOM on very large PDFs
                    if self.max_pages and page_num >= self.max_pages:
                        logger.info(
                            "PDF extraction capped at %d pages (max_pages=%d)",
                            page_num, self.max_pages,
                        )
                        break
                    try:
                        raw = page.extract_text() or ""
                        # Only extract tables for smaller pages to avoid OOM on dense layouts
                        table_txt = ""
                        if len(raw) < 10_000:
                            tables = page.extract_tables()
                            for tbl in tables:
                                for row in tbl:
                                    if row:
                                        table_txt += " | ".join(str(c or "") for c in row) + "\n"
                        combined = _clean_text(raw + "\n" + table_txt)
                    except Exception as page_err:
                        logger.warning("Skipping page %d: %s", page_num + 1, page_err)
                        combined = ""
                    finally:
                        # Release page object immediately to free pdfplumber's render buffers
                        del page
                        if page_num % 10 == 0:
                            gc.collect()

                    if combined:
                        pages_text.append(combined)
                        total_chars += len(combined)
                        if self.max_raw_chars and total_chars >= self.max_raw_chars:
                            logger.info(
                                "PDF extraction capped at %d chars after %d pages "
                                "(max_raw_chars=%d)",
                                total_chars, len(pages_text), self.max_raw_chars,
                            )
                            break
        except MemoryError:
            logger.error(
                "OOM while processing PDF '%s' after %d pages — partial result returned.",
                path.name, len(pages_text),
            )
            if not pages_text:
                raise RuntimeError(
                    f"Not enough memory to process '{path.name}'. "
                    "Try a smaller file or increase Docker memory allocation."
                )

        return pages_text, len(pages_text)

    def _extract_docx(
        self, path: Path, file_obj: Optional[IO[bytes]]
    ) -> tuple[List[str], int]:
        """Extract paragraph text from DOCX files."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        source = file_obj if file_obj is not None else str(path)
        doc = DocxDocument(source)

        paragraphs: List[str] = []
        current_page: List[str] = []

        total_chars = 0
        for para in doc.paragraphs:
            text = _clean_text(para.text)
            if text:
                current_page.append(text)
                total_chars += len(text)
            # Simulate page breaks every ~50 paragraphs
            if len(current_page) >= 50:
                paragraphs.append("\n".join(current_page))
                current_page = []
            if self.max_raw_chars and total_chars >= self.max_raw_chars:
                break

        if current_page:
            paragraphs.append("\n".join(current_page))

        return paragraphs or [""], max(1, len(paragraphs))

    def _extract_text(
        self, path: Path, file_obj: Optional[IO[bytes]]
    ) -> tuple[List[str], int]:
        """Read plain text / Markdown files."""
        if file_obj is not None:
            raw = file_obj.read().decode("utf-8", errors="replace")
        else:
            raw = path.read_text(encoding="utf-8", errors="replace")

        cleaned = _clean_text(raw)
        if self.max_raw_chars and len(cleaned) > self.max_raw_chars:
            cleaned = cleaned[: self.max_raw_chars]
        # Split on blank lines to create logical "pages"
        sections = [s.strip() for s in re.split(r"\n\n+", cleaned) if s.strip()]
        return sections or [cleaned], len(sections) or 1

    # ── Chunk Builder ─────────────────────────────────────────

    def _build_chunks(
        self, pages_text: List[str], doc_id: str, doc_name: str
    ) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        chunk_index = 0

        for page_num, page_text in enumerate(pages_text):
            raw_chunks = _chunk_text(page_text, self.chunk_size, self.overlap)
            for raw in raw_chunks:
                cid = _stable_id(f"{doc_id}:{chunk_index}:{raw[:50]}")
                chunks.append(
                    DocumentChunk(
                        chunk_id=cid,
                        doc_id=doc_id,
                        doc_name=doc_name,
                        page_num=page_num,
                        chunk_index=chunk_index,
                        text=raw,
                        metadata={
                            "source": doc_name,
                            "page": page_num + 1,
                            "chunk_index": chunk_index,
                        },
                    )
                )
                chunk_index += 1

        return chunks


# ── Processor factory ──────────────────────────────────────────────────────────

def get_processor(
    use_docling: bool = True,
    use_ocr: bool = False,
    chunk_size: int = 800,
    overlap: int = 150,
    max_raw_chars: int = 0,
    max_pages: int = 300,
):
    """
    Return a DocumentProcessor or DoclingProcessor based on settings.

    DoclingProcessor is returned when use_docling=True and provides advanced
    PDF parsing, table extraction, OCR, and support for PPTX/XLSX/HTML/images.
    Falls back to DocumentProcessor when use_docling=False (fast, no ML models).
    """
    if use_docling:
        from tools.docling_processor import DoclingProcessor
        return DoclingProcessor(use_ocr=use_ocr, max_raw_chars=max_raw_chars)
    return DocumentProcessor(
        chunk_size=chunk_size, overlap=overlap,
        max_raw_chars=max_raw_chars, max_pages=max_pages,
    )
